import os
import xlrd
import pandas as pd
from docx import Document
from docx.shared import Mm
from docx.shared import Inches

def load_excel(file_path):
	try:
		# Load proms data from Excel file
		proms = pd.read_excel(file_path)
		page_num = proms.shape[0]
		return proms, page_num
	except Exception as e:
		print(f"Error reading Excel file: {e}")
		return None

def load_pics(file_path):
	try:
		# Load proms pictures from folder
		pics = os.listdir(file_path)
		return pics
	except Exception as e:
		print(f"Error reading pictures folder: {e}")
		return None

def create_docx(proms, page_num, pics):
	# Create Word document
	document = Document()

	# Change page orientation with standard A4 format
	section = document.sections[0]
	section.page_height = Mm(210)
	section.page_width = Mm(297)

	for i in range(0, page_num - 3):
		content = ''

		# Add prom STYLE # and prom colors
		style = int(proms.loc[i].at['STYLE #']) # Get prom STYLE #
		content += str(style) + "\n"

		# Get prom colors
		prom = proms.loc[i]
		prom = prom.drop(labels=['Unnamed: 0', 'DEV #', 'STYLE #', 'FACTORY', 'COLOR REQUEST', 'NOTES', 'PRICE'])
		for j in prom:
			if pd.isna(j) == False:
				content += j + ", "

		content = content[:-1] + '\n\n' # Remove last element in string (',')

		paragraph = document.add_paragraph(content) # Insert prom STYLE # and prom colors into file
		paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify ....

		# Add prom pics
		r = paragraph.add_run()

		dev = proms.loc[i].at['DEV #'] # Get prom DEV #
		dev = dev.split('-')[1]

		for j in pics:
			tmp = j.split()
			if tmp[0] == dev:
				r.add_picture('./dresses on models/' + j, width=Inches(2.24), height=Inches(4))

		# Add blank page
		document.add_page_break()

		print("Row " + str(i + 1) + " Done..")

	document.save('demo.docx')

def main():
	excel_path = 'PROM.xlsx'
	pic_path = './dresses on models'

	proms, page_num = load_excel(excel_path)
	pics = load_pics(pic_path)

	if proms is not None:
		create_docx(proms, page_num, pics)
		print("DOCX file created successfully.")
	else:
		print("Failed to create DOCX file.")

if __name__ == "__main__":
	main()
